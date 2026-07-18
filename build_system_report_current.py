from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("NCUT_eduroam_RADIUS_system_report.docx").resolve()

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 96, 105)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
SUCCESS_FILL = "EAF6EA"
CAUTION_FILL = "FFF8E8"
BORDER = "C8D0DA"


def set_run_font(run, name="Calibri", east_asia="Microsoft JhengHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{margin}"))
        if el is None:
            el = OxmlElement(f"w:{margin}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=BLACK, size=9.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para(p, before=0, after=0, line=1.10)
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_widths(table, widths_dxa, indent=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index >= len(widths_dxa):
                continue
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_GRAY, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, widths_dxa)
    set_table_borders(table)
    keep_row_together(table.rows[0])
    repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, size=font_size)
    for row in rows:
        docx_row = table.add_row()
        keep_row_together(docx_row)
        for i, value in enumerate(row):
            set_cell_text(docx_row.cells[i], value, size=font_size)
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=6, line=1.0)
    return table


def add_kv_table(doc, rows, label_width=2300, value_width=7060):
    return add_table(doc, ["項目", "設定值"], rows, [label_width, value_width], font_size=9.2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        set_para(p, before=16, after=8)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        set_para(p, before=12, after=6)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        set_para(p, before=8, after=4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para(p)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_callout(doc, title, body, fill=CALLOUT, accent=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [9360])
    set_table_borders(table, color="D7DEE8")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""

    title_p = cell.paragraphs[0]
    set_para(title_p, before=0, after=3)
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=10.5, color=accent, bold=True)

    body_p = cell.add_paragraph()
    set_para(body_p, before=0, after=0)
    body_run = body_p.add_run(body)
    set_run_font(body_run, size=10.5, color=BLACK)
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=6)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [9360])
    set_table_borders(table, color="DADDE3")
    cell = table.cell(0, 0)
    shade_cell(cell, "FAFAFB")
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    cell.text = ""
    for index, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        set_para(p, before=0, after=0, line=1.0)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Microsoft JhengHei", size=8.2, color=BLACK)
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=4)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def paragraph_bottom_border(paragraph, color="A6B3C2", size="8", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def apply_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    doc.styles["Normal"].font.size = Pt(11)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        h = doc.styles[style_name]
        h.font.size = Pt(size)
        h.font.color.rgb = color
        h.font.bold = True
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.line_spacing = 1.10


def setup_header_footer(doc):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run("NCUT eduroam RADIUS 系統設定報告書")
    set_run_font(header_run, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Page ")
    set_run_font(footer_run, size=9, color=MUTED)
    add_page_number(footer_p)


CLIENT_ROWS = [
    ("TANRC_MainAuth_SRV", "10.1.77.0/24", "TANRC 主驗證網段"),
    ("TANRC_NODES_SRV", "10.1.78.0/24", "TANRC 節點網段"),
    ("10.1.77.7", "10.1.77.7", "roamingcenter"),
    ("10.1.77.11", "10.1.77.11", "roamingcenter-monitor"),
    ("140.128.71.11", "140.128.71.11", "ncut"),
    ("140.128.71.31", "140.128.71.31", "student"),
    ("140.128.72.34", "140.128.72.34", "fg620b"),
    ("140.128.72.6", "140.128.72.6", "fg620b"),
    ("192.168.255.2", "192.168.255.2", "FG1500D-vdom-wireless-0"),
    ("140.128.71.210", "140.128.71.210", "student"),
    ("140.128.71.213", "140.128.71.213", "student"),
    ("140.128.71.211", "140.128.71.211", "student"),
    ("140.128.71.212", "140.128.71.212", "student"),
    ("140.128.71.214", "140.128.71.214", "student"),
    ("140.128.71.215", "140.128.71.215", "student"),
    ("140.128.71.195", "140.128.71.195", "student"),
    ("140.128.77.130", "140.128.77.130", "fg3000"),
    ("140.128.71.251", "140.128.71.251", "vlan71"),
    ("140.128.72.134", "140.128.72.134", "fg1500D"),
    ("140.128.72.243", "140.128.72.243", "3810A"),
    ("120.108.14.167", "120.108.14.167", "TEST"),
    ("120.108.11.200", "120.108.11.200", "FG-61F"),
    ("120.108.30.253", "120.108.30.253", "FG-1800F"),
]


RADIUS_CLIENTS_SNIPPET = """client TANRC_MainAuth_SRV {
        ipaddr = 10.1.77.0/24
        secret = [REDACTED]
}

client TANRC_NODES_SRV {
        ipaddr = 10.1.78.0/24
        secret = [REDACTED]
}

# Additional NCUT NAS / controller clients imported from local clients.conf.
# 21 device client blocks are present. Shared secrets are not disclosed."""


RADIUS_PROXY_SNIPPET = """home_server TANRC_MainAuth_SRV1 {
        type = auth+acct
        ipaddr = 10.1.77.7
        port = 1812
        secret = [REDACTED]
        response_window = 3
        zombie_period = 180
        status_check = status-server
}

home_server TANRC_MainAuth_SRV2 {
        type = auth+acct
        ipaddr = 10.1.77.9
        port = 1812
        secret = [REDACTED]
        response_window = 3
        zombie_period = 180
        status_check = status-server
}

home_server_pool TANRC_POOL {
        type = fail-over
        home_server = TANRC_MainAuth_SRV1
        home_server = TANRC_MainAuth_SRV2
}

realm DEFAULT {
        pool = TANRC_POOL
        nostrip
}"""


OPENVPN_CONFIG_SNIPPET = """client
dev tun
proto tcp
remote 163.28.170.30 443
remote 163.28.192.141 443
resolv-retry infinite
nobind
persist-tun
remote-cert-tls server
cipher AES-256-GCM
auth SHA512
auth-nocache
tls-version-min 1.2
tls-crypt-v2 [KEY FILE REDACTED]
verb 3
log-append /var/log/openvpn/openvpn.log"""


OPENVPN_SUCCESS_LOG = """TCP connection established with [AF_INET]163.28.170.30:443
VERIFY OK: depth=1, CN=TANRC
VERIFY OK: depth=0, CN=MainVPN
Control Channel: TLSv1.3, cipher TLS_AES_256_GCM_SHA384
PUSH_REPLY ... ifconfig 10.1.0.59 255.255.0.0
TUN/TAP device tun0 opened
net_addr_v4_add: 10.1.0.59/16 dev tun0
Initialization Sequence Completed"""


def build_doc():
    doc = Document()
    apply_doc_styles(doc)
    setup_header_footer(doc)

    p = doc.add_paragraph()
    set_para(p, before=18, after=2)
    set_run_font(p.add_run("系統設定報告書"), size=12, color=MUTED, bold=True)

    p = doc.add_paragraph()
    set_para(p, before=0, after=4)
    set_run_font(p.add_run("國立勤益科技大學 eduroam RADIUS"), size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    set_para(p, before=0, after=14)
    set_run_font(p.add_run("Rocky Linux 9 / FreeRADIUS / OpenVPN"), size=13.5, color=MUTED)
    paragraph_bottom_border(p)

    add_kv_table(doc, [
        ("伺服器 IP", "140.128.71.249"),
        ("Hostname", "71-249"),
        ("作業系統", "Rocky Linux 9.8 (Blue Onyx)"),
        ("Kernel", "5.14.0-687.10.1.el9_8.0.1.x86_64"),
        ("報告產生時間", "2026-06-18 14:41 CST"),
        ("資料來源", "遠端主機即時檢查、建置手冊、clients.conf、OpenVPN client package"),
        ("機敏資訊處理", "shared secret、private key、tls-crypt-v2 key 均不揭露，以 [REDACTED] 表示。"),
    ])

    add_callout(
        doc,
        "整體狀態",
        "FreeRADIUS、firewalld、OpenVPN client 與 sshd 均為 active/enabled。OpenVPN 已成功建立 tun0，VPN IP 為 10.1.0.59/16；對 TANRC 10.1.77.7 與 10.1.77.9 的 RADIUS Status-Server 測試皆收到 Access-Accept。",
        fill=SUCCESS_FILL,
        accent=RGBColor(38, 103, 52),
    )

    add_heading(doc, "1. 執行摘要", 1)
    add_table(doc, ["項目", "目前狀態", "說明"], [
        ("FreeRADIUS", "正常", "radiusd active/enabled，radiusd -XC 顯示 Configuration appears to be OK。"),
        ("本機 RADIUS 測試", "通過", "127.0.0.1:1812 Status-Server 收到 Access-Accept。"),
        ("TANRC RADIUS 測試", "通過", "10.1.77.7 與 10.1.77.9 皆透過 tun0 回應 Access-Accept。"),
        ("OpenVPN", "正常", "openvpn-client@client active/enabled，Status 為 Initialization Sequence Completed。"),
        ("tun0", "已建立", "介面 tun0 位址為 10.1.0.59/16，路由 10.1.0.0/16 走 tun0。"),
        ("防火牆", "正常", "firewalld public zone 已開放 UDP 1812、1813。"),
        ("備援 VPN endpoint", "需追蹤", "163.28.192.141:443 仍回 Connection refused；主 endpoint 163.28.170.30:443 已可用。"),
    ], [1900, 1600, 5860], header_fill=BLUE_GRAY)

    add_heading(doc, "2. 系統與網路設定", 1)
    add_kv_table(doc, [
        ("主機名稱", "71-249"),
        ("對外來源 IP", "140.128.71.249"),
        ("主要介面", "ens192: 140.128.71.249/24"),
        ("IPv6", "2001:288:5007:2071:250:56ff:fe99:b30e/64"),
        ("VPN 介面", "tun0: 10.1.0.59/16"),
        ("預設閘道", "140.128.71.254 via ens192"),
        ("DNS", "140.128.71.1, 140.128.71.3"),
    ])
    add_table(doc, ["目的網段", "路由", "用途"], [
        ("default", "via 140.128.71.254 dev ens192", "一般對外連線與 SSH 管理。"),
        ("140.128.71.0/24", "dev ens192 src 140.128.71.249", "校內本機網段。"),
        ("10.1.0.0/16", "dev tun0 src 10.1.0.59", "TANRC VPN 內部網段，包含 10.1.77.7 與 10.1.77.9。"),
    ], [2200, 3800, 3360])

    add_heading(doc, "3. 套件、服務與主機防火牆", 1)
    add_table(doc, ["套件", "版本"], [
        ("freeradius", "3.0.27-3.el9_7.x86_64"),
        ("freeradius-utils", "3.0.27-3.el9_7.x86_64"),
        ("openvpn", "2.5.11-1.el9.x86_64"),
        ("firewalld", "1.3.4-18.el9_7.noarch"),
        ("epel-release", "9-10.el9.noarch"),
        ("net-tools", "2.0-0.64.20160912git.el9.x86_64"),
        ("tar", "1.34-11.el9.x86_64"),
    ], [3000, 6360])
    add_table(doc, ["服務", "Active", "Enabled", "說明"], [
        ("radiusd", "active", "enabled", "FreeRADIUS daemon。"),
        ("firewalld", "active", "enabled", "主機防火牆。"),
        ("openvpn-client@client", "active", "enabled", "TANRC OpenVPN client tunnel。"),
        ("sshd", "active", "enabled", "遠端管理服務。"),
    ], [2600, 1300, 1300, 4160])
    add_kv_table(doc, [
        ("firewalld state", "running"),
        ("active zone", "public, interface ens192"),
        ("public ports", "1812/udp 1813/udp"),
    ])

    add_heading(doc, "4. FreeRADIUS 設定", 1)
    add_body(doc, "FreeRADIUS 以 /etc/raddb/clients.conf 定義可送出 RADIUS request 的來源端，並以 /etc/raddb/proxy.conf 將 DEFAULT realm 轉送至 TANRC home server pool。設定檔中的 shared secret 已遮蔽。")
    add_table(doc, ["檢查項目", "結果"], [
        ("radiusd -XC", "Configuration appears to be OK"),
        ("client block 總數", "25"),
        ("NCUT clients.conf 匯入區塊", "存在；21 筆設備 client 已匯入。"),
        ("目前備份檔", "/etc/raddb/clients.conf.codex-backup-20260618-122351"),
    ], [3000, 6360])

    add_heading(doc, "4.1 RADIUS Client 清單", 2)
    add_table(doc, ["Client", "IP / CIDR", "Shortname / 用途"], CLIENT_ROWS, [2700, 2400, 4260], font_size=8.4)

    add_heading(doc, "4.2 Proxy 與 Realm", 2)
    add_table(doc, ["Home server", "IP", "Port", "Type", "Status check"], [
        ("TANRC_MainAuth_SRV1", "10.1.77.7", "1812", "auth+acct", "status-server"),
        ("TANRC_MainAuth_SRV2", "10.1.77.9", "1812", "auth+acct", "status-server"),
    ], [2600, 1700, 900, 1500, 2660])
    add_table(doc, ["Realm", "Pool", "轉送策略"], [
        ("DEFAULT", "TANRC_POOL", "fail-over pool，nostrip enabled"),
    ], [2200, 2500, 4660])

    add_heading(doc, "5. OpenVPN 設定與連線狀態", 1)
    add_table(doc, ["項目", "設定值"], [
        ("OpenVPN package", "openvpn-2.5.11-1.el9.x86_64"),
        ("服務", "openvpn-client@client active/enabled"),
        ("連線狀態", "Initialization Sequence Completed"),
        ("Client certificate", "subject CN=U_59_NCUT, issuer CN=TANRC"),
        ("憑證有效期", "2026-06-15 至 2036-06-12"),
        ("主 VPN endpoint", "163.28.170.30:443 CONNECT_OK"),
        ("備援 VPN endpoint", "163.28.192.141:443 CONNECT_FAIL / Connection refused"),
        ("VPN 介面", "tun0 10.1.0.59/16"),
        ("VPN 路由", "10.1.0.0/16 dev tun0 src 10.1.0.59"),
    ], [2600, 6760])
    add_table(doc, ["檔案", "權限", "說明"], [
        ("client.conf", "0644", "OpenVPN client configuration"),
        ("ca.crt", "0644", "TANRC CA certificate"),
        ("nU_59_NCUT.crt", "0644", "Client certificate"),
        ("nU_59_NCUT.key", "0600", "Client private key; content not disclosed"),
        ("ta.key", "0600", "Static key file present"),
        ("tc_client.key", "0600", "tls-crypt-v2 client key present"),
        ("nU_59_NCUT.tar", "0600", "Original VPN package retained under /etc/openvpn/client"),
    ], [2700, 1200, 5460])

    add_callout(
        doc,
        "VPN 測試結果",
        "防火牆更新後，OpenVPN 已完成 TLSv1.3 控制通道、取得 10.1.0.59/16，並建立 tun0。RADIUS 到 TANRC 10.1.77.7 與 10.1.77.9 的測試封包目前均走 tun0。",
        fill=SUCCESS_FILL,
        accent=RGBColor(38, 103, 52),
    )

    add_heading(doc, "6. 測試紀錄", 1)
    add_table(doc, ["測試項目", "方法 / 目標", "結果", "重點輸出"], [
        ("FreeRADIUS 語法檢查", "radiusd -XC", "通過", "Configuration appears to be OK"),
        ("本機 RADIUS 狀態測試", "radclient 127.0.0.1:1812 status", "通過", "Received Access-Accept"),
        ("TANRC 主節點狀態測試", "radclient 10.1.77.7:1812 status", "通過", "Received Access-Accept, route dev tun0 src 10.1.0.59"),
        ("TANRC 備援節點狀態測試", "radclient 10.1.77.9:1812 status", "通過", "Received Access-Accept, route dev tun0 src 10.1.0.59"),
        ("OpenVPN 主 endpoint", "TCP 163.28.170.30:443", "通過", "TCP CONNECT_OK, TLS handshake completed"),
        ("OpenVPN 備援 endpoint", "TCP 163.28.192.141:443", "待 TANRC 確認", "Connection refused"),
    ], [2100, 2500, 1500, 3260], header_fill=BLUE_GRAY, font_size=8.6)

    add_heading(doc, "7. 防火牆與網路規則建議", 1)
    add_table(doc, ["方向", "來源", "目的", "Port / Protocol", "用途 / 狀態"], [
        ("Outbound", "140.128.71.249", "163.28.170.30", "TCP/443", "OpenVPN 主節點；目前可連並已完成 VPN。"),
        ("Outbound", "140.128.71.249", "163.28.192.141", "TCP/443", "OpenVPN 備援節點；目前對端拒絕連線，需 TANRC 確認服務狀態。"),
        ("Inbound", "校內 WLAN controller / AP / NAS", "140.128.71.249", "UDP/1812", "RADIUS authentication。"),
        ("Inbound", "校內 WLAN controller / AP / NAS", "140.128.71.249", "UDP/1813", "RADIUS accounting。"),
        ("VPN internal", "10.1.0.59", "10.1.77.7 / 10.1.77.9", "UDP/1812", "TANRC Status-Server 與 proxy auth 測試已通過。"),
    ], [1200, 2200, 2200, 1500, 2260], header_fill=BLUE_GRAY, font_size=8.0)
    add_body(doc, "Outbound OpenVPN TCP/443 不應套用 SSL inspection、HTTP proxy 或 DPI 轉譯，避免 OpenVPN TLS control channel 無法完成握手。")

    add_heading(doc, "8. 待追蹤事項", 1)
    add_table(doc, ["項目", "目前狀態", "建議處理"], [
        ("OpenVPN 備援 endpoint", "163.28.192.141:443 回 Connection refused。", "請 TANRC 確認該節點是否啟用，或是否只提供主節點服務。"),
        ("校內 NAS 連線測試", "伺服器端已開 UDP 1812/1813。", "由 WLAN controller/AP/NAS 對 140.128.71.249 發送實際 auth/accounting 測試。"),
        ("報告維護", "本報告反映 2026-06-18 14:41 CST 狀態。", "若後續調整 clients.conf、proxy.conf 或 VPN endpoint，應同步更新報告與變更紀錄。"),
    ], [2300, 3300, 3760], header_fill=CAUTION_FILL, font_size=8.8)

    doc.add_page_break()
    add_heading(doc, "附錄 A. 設定摘要", 1)
    add_heading(doc, "A.1 /etc/raddb/clients.conf", 2)
    add_code_block(doc, RADIUS_CLIENTS_SNIPPET)
    add_heading(doc, "A.2 /etc/raddb/proxy.conf", 2)
    add_code_block(doc, RADIUS_PROXY_SNIPPET)
    doc.add_page_break()
    add_heading(doc, "A.3 /etc/openvpn/client/client.conf 摘要", 2)
    add_code_block(doc, OPENVPN_CONFIG_SNIPPET)
    add_heading(doc, "A.4 OpenVPN 成功連線 log 摘要", 2)
    add_code_block(doc, OPENVPN_SUCCESS_LOG)

    doc.core_properties.title = "NCUT eduroam RADIUS 系統設定報告書"
    doc.core_properties.subject = "Rocky Linux 9 FreeRADIUS OpenVPN configuration report"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Preset: standard_business_brief; first-page header pattern: memo_masthead; secrets redacted."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
