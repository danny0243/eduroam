from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("NCUT_eduroam_RADIUS_system_report_20260616.docx").resolve()

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 96, 105)
BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
BORDER = "C8D0DA"


def set_run_font(run, name="Calibri", east_asia="Microsoft JhengHei", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para(p, before=0, after=6, line=1.10):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = tc_mar.find(qn(f"w:{margin}"))
        if el is None:
            el = OxmlElement(f"w:{margin}")
            tc_mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=BLACK, size=9.3, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para(p, before=0, after=0, line=1.10)
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_widths(table, widths_dxa, indent=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index >= len(widths_dxa):
                continue
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_GRAY, font_size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, widths_dxa)
    set_table_borders(table)
    keep_row_together(table.rows[0])
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, size=font_size)
    for row in rows:
        docx_row = table.add_row()
        keep_row_together(docx_row)
        cells = docx_row.cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=6, line=1.0)
    return table


def add_kv_table(doc, rows):
    return add_table(doc, ["欄位", "值"], rows, [2600, 6760], font_size=9.3)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        set_para(p, before=16, after=8)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        set_para(p, before=12, after=6)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        set_para(p, before=8, after=4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_para(p)
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    return p


def add_callout(doc, title, body, fill=CALLOUT, accent=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [9360])
    set_table_borders(table, color="D7DEE8")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""

    title_p = cell.paragraphs[0]
    set_para(title_p, before=0, after=3)
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=10.5, color=accent, bold=True)

    body_p = cell.add_paragraph()
    set_para(body_p, before=0, after=0)
    body_run = body_p.add_run(body)
    set_run_font(body_run, size=10.5, color=BLACK)
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=6)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [9360])
    set_table_borders(table, color="DADDE3")
    cell = table.cell(0, 0)
    shade_cell(cell, "FAFAFB")
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    cell.text = ""
    for index, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        set_para(p, before=0, after=0, line=1.0)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", east_asia="Microsoft JhengHei", size=8.4, color=BLACK)
    spacer = doc.add_paragraph()
    set_para(spacer, before=0, after=4)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def apply_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    doc.styles["Normal"].font.size = Pt(11)
    for style_name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)):
        doc.styles[style_name].font.size = Pt(size)
        doc.styles[style_name].font.color.rgb = color
        doc.styles[style_name].font.bold = True


def setup_header_footer(doc):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run("NCUT eduroam RADIUS 系統設定報告書")
    set_run_font(header_run, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Page ")
    set_run_font(footer_run, size=9, color=MUTED)
    add_page_number(footer_p)


RADIUS_CLIENTS_SNIPPET = """# BEGIN NCUT TANRC CLIENTS - managed by Codex 2026-06-16
client TANRC_MainAuth_SRV {
        ipaddr = 10.1.77.0/24
        secret = [REDACTED]
}

client TANRC_NODES_SRV {
        ipaddr = 10.1.78.0/24
        secret = [REDACTED]
}
# END NCUT TANRC CLIENTS"""

RADIUS_PROXY_SNIPPET = """# BEGIN NCUT TANRC PROXY - managed by Codex 2026-06-16
home_server TANRC_MainAuth_SRV1 {
        type = auth+acct
        ipaddr = 10.1.77.7
        port = 1812
        secret = [REDACTED]
        response_window = 3
        zombie_period = 180
        status_check = status-server
}

home_server TANRC_MainAuth_SRV2 {
        type = auth+acct
        ipaddr = 10.1.77.9
        port = 1812
        secret = [REDACTED]
        response_window = 3
        zombie_period = 180
        status_check = status-server
}

home_server_pool TANRC_POOL {
        type = fail-over
        home_server = TANRC_MainAuth_SRV1
        home_server = TANRC_MainAuth_SRV2
}

realm DEFAULT {
        pool = TANRC_POOL
        nostrip
}
# END NCUT TANRC PROXY"""

OPENVPN_CONFIG_SNIPPET = """client
dev tun
proto tcp
remote 163.28.170.30 443
remote 163.28.192.141 443
resolv-retry infinite
nobind
persist-tun
remote-cert-tls server
cipher AES-256-GCM
auth SHA512
auth-nocache
tls-version-min 1.2
tls-crypt-v2 [KEY FILE]
verb 3
log-append /var/log/openvpn/openvpn.log"""


def build_doc():
    doc = Document()
    apply_doc_styles(doc)
    setup_header_footer(doc)

    p = doc.add_paragraph()
    set_para(p, before=18, after=2)
    set_run_font(p.add_run("系統設定報告書"), size=12, color=MUTED, bold=True)

    p = doc.add_paragraph()
    set_para(p, before=0, after=4)
    set_run_font(p.add_run("eduroam RADIUS 服務建置與現況"), size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    set_para(p, before=0, after=14)
    set_run_font(p.add_run("Rocky Linux 9 / FreeRADIUS / OpenVPN"), size=13.5, color=MUTED)

    add_kv_table(doc, [
        ("目標主機", "140.128.71.249"),
        ("主機名稱", "71-249"),
        ("作業系統", "Rocky Linux 9.8"),
        ("資料擷取時間", "2026-06-16 20:22:00 CST"),
        ("報告依據", "建置手冊_Rocky Linux 9(General).docx 與實機查核結果"),
        ("敏感資訊處理", "帳號密碼、shared secret、private key 內容均未列入；設定片段以 [REDACTED] 標示。"),
    ])

    add_callout(
        doc,
        "整體結論",
        "FreeRADIUS 服務已完成安裝、設定、啟用與本機 RADIUS 狀態測試；OpenVPN 用戶端與金鑰已部署並設為開機啟用，但 VPN tunnel 尚未建立，故對 TANRC 10.1.77.7 / 10.1.77.9 的漫遊測試仍無法取得 Access-Accept。",
    )

    add_heading(doc, "1. 摘要狀態", 1)
    add_table(doc, ["項目", "目前狀態", "佐證"], [
        ("FreeRADIUS", "已完成 / 正常", "radiusd active/enabled；radiusd -XC 顯示 Configuration appears to be OK。"),
        ("本機 RADIUS 測試", "通過", "127.0.0.1:1812 Status-Server 收到 Access-Accept。"),
        ("主機防火牆", "已開通", "firewalld public zone 已永久開啟 UDP 1812、1813。"),
        ("OpenVPN 用戶端", "服務已啟動 / tunnel 未建立", "openvpn-client@client active/enabled，但 tun0_absent=1。"),
        ("TANRC 漫遊測試", "尚未通過", "local radtest 經 proxy 後因 home server 無回應而 Access-Reject；direct radtest timeout。"),
    ], [1900, 2100, 5360], header_fill=BLUE_GRAY)

    add_heading(doc, "2. 主機與網路環境", 1)
    add_kv_table(doc, [
        ("主機 IP", "140.128.71.249"),
        ("主機名稱", "71-249"),
        ("作業系統", "Rocky Linux 9.8"),
        ("Kernel", "5.14.0-687.10.1.el9_8.0.1.x86_64"),
        ("對外來源 IP", "140.128.71.249"),
    ])
    add_table(doc, ["介面", "狀態", "位址"], [
        ("lo", "UNKNOWN", "127.0.0.1/8, ::1/128"),
        ("ens192", "UP", "140.128.71.249/24, 2001:288:5007:2071:250:56ff:fe99:b30e/64, fe80::250:56ff:fe99:b30e/64"),
    ], [1400, 1500, 6460])
    add_body(doc, "目前到 10.1.77.7 與 10.1.77.9 的路由仍經由 140.128.71.254 / ens192，表示 OpenVPN tunnel 尚未接管 TANRC 內部網段。")

    add_heading(doc, "3. 套件與服務", 1)
    add_table(doc, ["套件", "版本 / 查核結果", "狀態"], [
        ("freeradius", "3.0.27-3.el9_7", "Installed"),
        ("freeradius-utils", "3.0.27-3.el9_7", "Installed"),
        ("openvpn", "2.5.11-1.el9", "Installed"),
        ("firewalld", "1.3.4-18.el9_7", "Installed"),
        ("epel-release", "9-10.el9", "Installed"),
        ("net-tools", "2.0-0.64.20160912git.el9", "Installed"),
        ("tar", "1.34-11.el9", "Installed"),
        ("vim package name", "rpm -q vim: not installed", "註：核心需求不依賴此套件名稱。"),
    ], [2500, 4300, 2560])
    add_table(doc, ["服務", "Active", "Enabled", "說明"], [
        ("radiusd", "active", "enabled", "FreeRADIUS service is running under systemd."),
        ("firewalld", "active", "enabled", "Host firewall is running."),
        ("openvpn-client@client", "active", "enabled", "OpenVPN process is running, but tunnel is not established."),
    ], [2500, 1500, 1500, 3860])

    add_heading(doc, "4. FreeRADIUS 設定", 1)
    add_heading(doc, "4.1 Client 設定", 2)
    add_table(doc, ["Client 名稱", "允許來源", "說明"], [
        ("TANRC_MainAuth_SRV", "10.1.77.0/24", "已設定；shared secret 已遮蔽"),
        ("TANRC_NODES_SRV", "10.1.78.0/24", "已設定；shared secret 已遮蔽"),
    ], [2600, 2200, 4560])
    add_heading(doc, "4.2 Proxy / Realm 設定", 2)
    add_table(doc, ["Home server", "IP", "Port", "Type", "角色"], [
        ("TANRC_MainAuth_SRV1", "10.1.77.7", "1812", "auth+acct", "Primary home server"),
        ("TANRC_MainAuth_SRV2", "10.1.77.9", "1812", "auth+acct", "Failover home server"),
    ], [2600, 1700, 900, 1200, 2960])
    add_body(doc, "realm DEFAULT 已設定為使用 TANRC_POOL，pool 類型為 fail-over，並保留 nostrip，使帳號 realm 不被移除。")
    doc.add_page_break()
    add_heading(doc, "4.3 監聽與防火牆", 2)
    add_table(doc, ["項目", "設定 / 結果"], [
        ("radiusd 監聽", "0.0.0.0:1812、0.0.0.0:1813、[::]:1812、[::]:1813"),
        ("firewalld public ports", "1812/udp 1813/udp"),
        ("設定備份", "/etc/raddb/clients.conf.codex-backup-20260616-142018；/etc/raddb/proxy.conf.codex-backup-20260616-142018"),
    ], [2600, 6760])

    add_heading(doc, "5. OpenVPN 設定", 1)
    add_table(doc, ["項目", "設定值"], [
        ("金鑰包", "nU_59_NCUT.tar"),
        ("部署目錄", "/etc/openvpn/client"),
        ("Client certificate", "CN=U_59_NCUT；Issuer=CN=TANRC；有效期 2026-06-15 至 2036-06-12"),
        ("協定 / 裝置", "proto tcp；dev tun"),
        ("遠端端點", "163.28.170.30:443；163.28.192.141:443"),
        ("加密設定", "cipher AES-256-GCM；auth SHA512；tls-version-min 1.2；tls-crypt-v2"),
        ("Log", "/var/log/openvpn/openvpn.log"),
    ], [2600, 6760])
    add_table(doc, ["檔案", "權限", "用途"], [
        ("client.conf", "0644", "OpenVPN client configuration"),
        ("ca.crt", "0644", "TANRC CA certificate"),
        ("nU_59_NCUT.crt", "0644", "Client certificate CN=U_59_NCUT"),
        ("nU_59_NCUT.key", "0600", "Client private key; content not disclosed"),
        ("ta.key", "0600", "Static key file present"),
        ("tc_client.key", "0600", "tls-crypt-v2 client key present"),
        ("nU_59_NCUT.tar", "0600", "Original key package retained in /etc/openvpn/client"),
    ], [2700, 1200, 5460])
    add_callout(
        doc,
        "OpenVPN 目前狀態",
        "openvpn-client@client 服務為 active/enabled，/dev/net/tun 存在，但 tun0 尚未產生。日誌顯示第一端點可建立 TCP 連線但 TLS key negotiation timeout；第二端點 TCP 443 connection refused。",
        fill="FFF8E8",
        accent=RGBColor(122, 90, 0),
    )

    add_heading(doc, "6. 測試結果", 1)
    add_table(doc, ["測試項目", "指令 / 方法", "結果", "判讀"], [
        ("FreeRADIUS 設定檢查", "radiusd -XC", "通過", "Configuration appears to be OK；exit status 0。"),
        ("本機 RADIUS 狀態測試", "radclient 127.0.0.1:1812 status testing123", "通過", "收到 Access-Accept。"),
        ("本機 proxy 漫遊測試", "radtest exmaple@rc.edu.tw -> 127.0.0.1", "未通過", "請求已由本機 RADIUS 接收，但因 TANRC home server 無回應而 Access-Reject。"),
        ("直接 TANRC 測試", "radtest -> 10.1.77.7", "未通過", "timeout；目前沒有 VPN tunnel 路由。"),
        ("VPN endpoint 1", "TCP connect 163.28.170.30:443", "TCP 可連", "OpenVPN TLS negotiation 逾時。"),
        ("VPN endpoint 2", "TCP connect 163.28.192.141:443", "拒絕連線", "Connection refused。"),
    ], [1900, 2600, 1300, 3560], header_fill=BLUE_GRAY, font_size=8.8)

    doc.add_page_break()
    add_heading(doc, "7. 未完成項目與原因", 1)
    add_body(doc, "RADIUS 服務本身已完成並通過本機狀態測試；目前未完成的是 TANRC VPN 連線建立與跨 VPN 的漫遊驗證。由於 tun0 不存在，10.1.77.7 / 10.1.77.9 仍走一般校園網路路由，RADIUS proxy 無法收到 TANRC home server 回應。")
    add_table(doc, ["風險 / 阻礙", "目前證據", "建議處理"], [
        ("VPN 第一端點 TLS 逾時", "163.28.170.30:443 TCP 可連，但 OpenVPN TLS key negotiation failed。", "請確認校園防火牆是否對 TCP 443 做 SSL inspection / proxy / DPI，並確認 TANRC 端是否允許來源 IP 140.128.71.249 與 U_59_NCUT 憑證。"),
        ("VPN 第二端點拒絕連線", "163.28.192.141:443 Connection refused。", "請向 TANRC 或網路管理單位確認該端點目前是否服務中、是否需改用其他備援端點。"),
        ("TANRC RADIUS 無法驗證", "direct radtest timeout；local proxy test 因 home server 無回應而 reject。", "待 OpenVPN tunnel 建立並取得 tun0 後重跑 radtest。"),
    ], [2300, 3300, 3760], header_fill="FFF8E8", font_size=8.8)

    add_heading(doc, "8. 防火牆開通需求", 1)
    add_table(doc, ["方向", "來源", "目的", "Port / Protocol", "用途 / 備註"], [
        ("Outbound", "140.128.71.249", "163.28.170.30", "TCP/443", "OpenVPN over TCP 443；需允許 established/related 回程，且不要做 SSL inspection / proxy / DPI。"),
        ("Outbound", "140.128.71.249", "163.28.192.141", "TCP/443", "OpenVPN 備援端點；目前遠端回 Connection refused，仍建議先保留規則。"),
        ("Inbound", "校內無線控制器 / AP / NAS 管理 IP", "140.128.71.249", "UDP/1812", "RADIUS authentication；來源建議限縮為校內設備。"),
        ("Inbound", "校內無線控制器 / AP / NAS 管理 IP", "140.128.71.249", "UDP/1813", "RADIUS accounting；來源建議限縮為校內設備。"),
    ], [1350, 2100, 2100, 1500, 2310], header_fill=BLUE_GRAY, font_size=8.2)
    add_body(doc, "若防火牆或路由設備會檢查 VPN 內側流量，另需允許本機 VPN IP 與 10.1.77.7 / 10.1.77.9 間的 UDP 1812、1813，以及 10.1.77.0/24、10.1.78.0/24 對本機 VPN IP 的 RADIUS 流量。一般校園邊界防火牆通常只會看到外層 TCP 443。")

    add_heading(doc, "9. 後續驗收步驟", 1)
    add_table(doc, ["順序", "動作", "完成條件"], [
        ("1", "網路管理單位開通 / 排除 TCP 443 攔截，並確認 TANRC 端金鑰與來源 IP 授權。", "OpenVPN log 出現 Initialization Sequence Completed。"),
        ("2", "確認 VPN tunnel。", "ip addr show tun0 可看到 10.1.x.x 位址；10.1.77.7 / 10.1.77.9 路由走 tun0。"),
        ("3", "重新執行手冊測試。", "radtest exmaple@rc.edu.tw exmaple123 127.0.0.1 0 testing123 收到 Access-Accept。"),
        ("4", "重新執行直接 TANRC 測試。", "radtest exmaple@rc.edu.tw exmaple123 10.1.77.7 0 [shared secret] 收到 Access-Accept。"),
        ("5", "與校內無線控制器 / AP / NAS 串接。", "校內設備可對 140.128.71.249 UDP 1812/1813 成功送出驗證與計費封包。"),
    ], [900, 4860, 3600])

    add_heading(doc, "附錄 A. 遮蔽後設定片段", 1)
    add_heading(doc, "A.1 /etc/raddb/clients.conf", 2)
    add_code_block(doc, RADIUS_CLIENTS_SNIPPET)
    add_heading(doc, "A.2 /etc/raddb/proxy.conf", 2)
    add_code_block(doc, RADIUS_PROXY_SNIPPET)
    add_heading(doc, "A.3 /etc/openvpn/client/client.conf 重要參數", 2)
    add_code_block(doc, OPENVPN_CONFIG_SNIPPET)
    add_heading(doc, "A.4 OpenVPN log 摘要", 2)
    add_code_block(doc, """163.28.170.30:443 TCP connection established, then TLS key negotiation failed within 60 seconds.
163.28.192.141:443 returned Connection refused.
No Initialization Sequence Completed message was observed, and tun0 is absent.""")

    doc.core_properties.title = "NCUT eduroam RADIUS 系統設定報告書"
    doc.core_properties.subject = "Rocky Linux 9 FreeRADIUS OpenVPN configuration report"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Preset: standard_business_brief; first-page header pattern: memo_masthead; secrets redacted."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
