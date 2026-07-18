from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("NCUT_eduroam_RADIUS_technical_guide.docx").resolve()
CAPTURE = Path("radius_technical_capture_20260625.txt")

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B667A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "AAB7C4"
WARN_FILL = "FFF4CC"
OK_FILL = "EAF4EF"


def parse_capture(path: Path) -> dict:
    data: dict[str, str] = {}
    sections: dict[str, list[str]] = {}
    if not path.exists():
        return data
    current = None
    for raw in path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw.rstrip("\n")
        if line.endswith("_BEGIN"):
            current = line[:-6]
            sections[current] = []
            continue
        if line.endswith("_END"):
            current = None
            continue
        if current:
            sections[current].append(line)
            continue
        if "=" in line:
            k, v = line.split("=", 1)
            data[k] = v
    for k, lines in sections.items():
        data[k] = "\n".join(lines)
    return data


def set_cell_text(cell, text, bold=False, color=INK, size=9.0, font="Calibri"):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa, indent=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_row_together(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            p_pr = p._p.get_or_add_pPr()
            keep = p_pr.find(qn("w:keepLines"))
            if keep is None:
                keep = OxmlElement("w:keepLines")
                p_pr.append(keep)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, widths_dxa)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    keep_row_together(table.rows[0])
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, size=font_size)
    for row in rows:
        docx_row = table.add_row()
        keep_row_together(docx_row)
        for i, value in enumerate(row):
            set_cell_text(docx_row.cells[i], value, size=font_size)
    doc.add_paragraph()
    return table


def add_kv_table(doc, rows, label_width=2700, value_width=6660):
    return add_table(doc, ["項目", "設定值"], rows, [label_width, value_width], font_size=9.1)


def add_status_table(doc, rows):
    return add_table(doc, ["檢查項目", "目前狀態", "維運判讀"], rows, [2500, 2500, 4360], font_size=8.8)


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [9360])
    set_table_borders(table, color="D7DEE8")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    for idx, line in enumerate(body.split("\n")):
        if idx:
            p2.add_run().add_break()
        run = p2.add_run(line)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def add_code(doc, text, label=None):
    if label:
        p_label = doc.add_paragraph()
        p_label.style = "CaptionSmall"
        p_label.add_run(label).bold = True
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_widths(table, [9360])
    set_table_borders(table, color="DADDE3")
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FB")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    cell.text = ""
    for idx, line in enumerate(text.strip("\n").split("\n")):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor.from_string("1F2937")
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def _next_id(existing, attr_name):
    vals = []
    for el in existing:
        val = el.get(qn(attr_name))
        if val is not None and str(val).isdigit():
            vals.append(int(val))
    return (max(vals) + 1) if vals else 1


def new_decimal_num(doc):
    numbering = doc.part.numbering_part.element
    abstract_id = _next_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    num_id = _next_id(numbering.findall(qn("w:num")), "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, val in [
        ("w:start", "1"),
        ("w:numFmt", "decimal"),
        ("w:lvlText", "%1."),
        ("w:lvlJc", "left"),
    ]:
        node = OxmlElement(tag)
        if tag == "w:numFmt":
            node.set(qn("w:val"), val)
        elif tag == "w:lvlText":
            node.set(qn("w:val"), val)
        elif tag == "w:lvlJc":
            node.set(qn("w:val"), val)
        else:
            node.set(qn("w:val"), val)
        lvl.append(node)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(tabs)
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def add_steps(doc, steps):
    num_id = new_decimal_num(doc)
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        apply_num(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(step)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def set_hyperlink_style(document):
    styles = document.styles
    if "CaptionSmall" not in styles:
        st = styles.add_style("CaptionSmall", 1)
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        st.font.size = Pt(9)
        st.font.color.rgb = RGBColor.from_string(MUTED)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        st.font.size = Pt(11)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.25

    set_hyperlink_style(doc)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = "NCUT eduroam RADIUS 技術說明文件"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p.runs:
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "文件內未保存 shared secret 或使用者明文密碼；實際機敏值請依校內密碼保管流程存放。"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fp.runs:
        fp.runs[0].font.size = Pt(8)
        fp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def title_page(doc, cap):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("NCUT eduroam RADIUS 技術說明文件")
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(INK)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(16)
    r = sp.add_run("Rocky Linux 9 / FreeRADIUS / TANRC roaming / SSL 憑證與帳號維運 SOP")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        doc,
        "文件用途",
        "本文件供國立勤益科技大學 eduroam RADIUS 服務日常維運、交接與稽核使用。"
        "\n內容涵蓋目前系統狀態、主要設定、帳號新增/修改流程、SSL 憑證更新流程、測試與故障排除。",
        fill=OK_FILL,
    )

    add_kv_table(
        doc,
        [
            ("產製日期", "2026-06-25"),
            ("資料基準時間", cap.get("CAPTURED_AT", "2026-06-25 15:31:25 CST")),
            ("主機名稱", cap.get("HOSTNAME", "71-249")),
            ("服務 FQDN", "eduroam.ncut.edu.tw"),
            ("服務 IP", "140.128.71.249"),
            ("文件敏感資訊策略", "密碼與 shared secret 一律以 [REDACTED] 或佔位符呈現。"),
        ],
    )

    add_heading(doc, "快速結論", 1)
    add_status_table(
        doc,
        [
            ("DNS", "eduroam.ncut.edu.tw -> 140.128.71.249", "用戶端可用此名稱作為 RADIUS Server Name。"),
            ("RADIUS 服務", "active / enabled", "FreeRADIUS 已開機自動啟動。"),
            ("Firewall", "1812/udp, 1813/udp", "RADIUS authentication/accounting 對外開放。"),
            ("VPN", "tun0: 10.1.0.59/16", "TANRC home server 10.1.77.7、10.1.77.9 經 tun0。"),
            ("SSL 憑證", "*.ncut.edu.tw，有效至 2026-11-05", "涵蓋 eduroam.ncut.edu.tw。"),
            ("本校 realm", "ncut.edu.tw、eduroam.ncut.edu.tw", "本機處理；外校 realm 仍走 TANRC_POOL。"),
        ],
    )

    doc.add_page_break()


def system_overview(doc, cap):
    add_heading(doc, "1. 系統架構與服務範圍", 1)
    p = doc.add_paragraph()
    p.add_run(
        "本機 71-249 承擔 NCUT eduroam RADIUS 服務。校內無線控制器與 TANRC 節點可向本機送出 RADIUS authentication/accounting；"
        "本校 realm 由本機 FreeRADIUS 處理，非本校 realm 則透過 OpenVPN tunnel 轉送至 TANRC home server。"
    )

    add_kv_table(
        doc,
        [
            ("作業系統", cap.get("OS", "Rocky Linux 9.8 (Blue Onyx)")),
            ("RADIUS 版本", cap.get("RADIUSD", "FreeRADIUS 3.0.27")),
            ("Hostname", cap.get("HOSTNAME", "71-249")),
            ("IPv4", cap.get("IPV4", "140.128.71.249/24 ens192; 10.1.0.59/16 tun0")),
            ("Default gateway", cap.get("DEFAULT_ROUTE", "default via 140.128.71.254 dev ens192")),
            ("DNS", f"eduroam.ncut.edu.tw -> {cap.get('DNS_EDUROAM', '140.128.71.249')}"),
        ],
    )

    add_heading(doc, "1.1 認證流程", 2)
    add_table(
        doc,
        ["情境", "處理流程", "目前設定"],
        [
            (
                "本校 eduroam 使用者",
                "使用者以 user@ncut.edu.tw 連線，FreeRADIUS 在本機 realm 處理，inner-tunnel 使用 MSCHAPv2 驗證本機帳號檔。",
                "realm ncut.edu.tw { nostrip }；users 檔帳號搭配 Proxy-To-Realm := LOCAL。",
            ),
            (
                "外校 roaming 使用者",
                "非本校 realm 由 realm DEFAULT 送往 TANRC_POOL，再由 TANRC 判斷上游機構。",
                "TANRC_MainAuth_SRV1=10.1.77.7；TANRC_MainAuth_SRV2=10.1.77.9；pool fail-over。",
            ),
            (
                "外部測試帳號",
                "ncut@rc.edu.tw 測試會走 DEFAULT -> TANRC_POOL。",
                "最後驗證結果為 Access-Accept，代表 roaming proxy 路徑可用。",
            ),
        ],
        [1900, 4600, 2860],
        font_size=8.6,
    )

    add_callout(
        doc,
        "重要 realm 注意事項",
        "目前明確本機處理的 realm 是 ncut.edu.tw 與 eduroam.ncut.edu.tw。"
        "\n若未來要讓 student.ncut.edu.tw、staff.ncut.edu.tw 等子網域也由本機處理，必須在 /etc/raddb/proxy.conf 另行新增對應 realm；否則會落入 DEFAULT 並轉送 TANRC。",
        fill=WARN_FILL,
    )


def config_inventory(doc, cap):
    add_heading(doc, "2. 主要設定檔與目錄", 1)
    add_table(
        doc,
        ["路徑", "用途", "維運重點"],
        [
            ("/etc/raddb/clients.conf", "RADIUS client allowlist 與 shared secret。", "新增控制器、TANRC 節點或測試來源時修改；shared secret 不可寫入報告。"),
            ("/etc/raddb/proxy.conf", "realm、home server、home_server_pool。", "ncut.edu.tw 本機處理；DEFAULT 轉送 TANRC_POOL。"),
            ("/etc/raddb/mods-config/files/authorize", "本機帳號密碼資料庫。", "新增/修改使用者帳號時修改；修改後跑 radiusd -XC 並重啟。"),
            ("/etc/raddb/users", "authorize 的相容性入口。", "目前為指向 mods-config/files/authorize 的使用方式。"),
            ("/etc/raddb/mods-enabled/eap", "EAP/PEAP/TLS 憑證與 inner-tunnel 設定。", "SSL 憑證更新時確認 private_key_file、certificate_file、ca_file。"),
            ("/etc/raddb/certs", "FreeRADIUS 使用的憑證與私鑰目錄。", "檔案建議 root:radiusd、0640；SELinux 環境需 restorecon。"),
            ("/var/log/radius/radius.log", "FreeRADIUS 主要 log。", "查 Login incorrect、Home Server failed、EAP session 等訊息。"),
            ("/etc/openvpn/client", "OpenVPN client 設定目錄。", "tun0 消失或 TANRC 不通時優先查 openvpn-client@client.service。"),
        ],
        [3300, 3000, 3060],
        font_size=8.5,
    )

    add_heading(doc, "2.1 目前服務狀態", 2)
    add_status_table(
        doc,
        [
            ("radiusd", f"{cap.get('RADIUSD_STATE', 'active')} / {cap.get('RADIUSD_ENABLED', 'enabled')}", "RADIUS authentication/accounting 主服務。"),
            ("firewalld", f"{cap.get('FIREWALLD_STATE', 'active')} / {cap.get('FIREWALLD_ENABLED', 'enabled')}", f"開放 port: {cap.get('FIREWALL_PORTS', '1812/udp 1813/udp')}"),
            ("OpenVPN", cap.get("OPENVPN_UNITS", "openvpn-client@client.service:active:running"), "提供 TANRC 10.1.77.0/24 路徑。"),
            ("tun0", cap.get("VPN_TUN0", "10.1.0.59/16"), "若 absent，TANRC proxy 大多會失敗。"),
        ],
    )

    add_heading(doc, "2.2 Client 與 Home Server 摘要", 2)
    clients = cap.get("CLIENTS_SUMMARY", "").splitlines()
    rows = []
    for line in clients:
        if "client " in line:
            line_no, rest = line.split(":", 1)
            name = rest.replace("client", "").replace("{", "").strip()
            rows.append((line_no, name, "詳情含 secret，請於伺服器上查閱 /etc/raddb/clients.conf。"))
    if rows:
        add_table(doc, ["行號", "Client", "說明"], rows, [900, 3200, 5260], font_size=8.3)

    add_table(
        doc,
        ["TANRC 項目", "設定", "用途"],
        [
            ("TANRC_MainAuth_SRV1", "10.1.77.7:1812", "主要 authentication/accounting home server。"),
            ("TANRC_MainAuth_SRV2", "10.1.77.9:1812", "備援 authentication/accounting home server。"),
            ("TANRC_POOL", "fail-over", "DEFAULT realm 轉送池。"),
            ("路由", cap.get("ROUTE_TANRC", "10.1.77.7 dev tun0"), "TANRC 經 OpenVPN tun0。"),
        ],
        [2400, 2600, 4360],
        font_size=8.6,
    )


def current_radius_config(doc, cap):
    add_heading(doc, "3. FreeRADIUS 目前關鍵設定", 1)

    add_heading(doc, "3.1 Realm 與 Proxy", 2)
    add_code(
        doc,
        """
realm ncut.edu.tw {
        nostrip
}

realm eduroam.ncut.edu.tw {
        nostrip
}

realm DEFAULT {
        pool = TANRC_POOL
        nostrip
}
""",
        "目前 proxy.conf 的核心 realm 邏輯",
    )
    add_bullets(
        doc,
        [
            "本校 user@ncut.edu.tw 不再落入 DEFAULT，因此不會被送往 TANRC 驗證。",
            "外校 realm 仍由 DEFAULT 轉送 TANRC_POOL，不影響 roaming。",
            "若出現 user@student.ncut.edu.tw 這類帳號，需視政策決定是否新增 student.ncut.edu.tw 為 local realm。",
        ]
    )

    add_heading(doc, "3.2 EAP / PEAP / TLS", 2)
    add_table(
        doc,
        ["項目", "目前設定", "說明"],
        [
            ("PEAP inner method", "mschapv2", "Windows、iOS、Android 常見 eduroam 設定。"),
            ("virtual_server", "inner-tunnel", "EAP 內層身份驗證走 inner-tunnel。"),
            ("TLS 版本", "min=1.2, max=1.2", "目前限制為 TLS 1.2。"),
            ("private_key_file", "${certdir}/ncutserver.key", "RADIUS server private key。"),
            ("certificate_file", "${certdir}/ncutserver-fullchain.pem", "server certificate + intermediate CA。"),
            ("ca_file", "${certdir}/ncutserver-uca.cer", "TWCA intermediate certificate。"),
        ],
        [2300, 3200, 3860],
        font_size=8.7,
    )

    add_heading(doc, "3.3 目前 SSL 憑證", 2)
    add_kv_table(
        doc,
        [
            ("Subject", "C=TW, ST=Taiwan, L=Taichung, O=National Chin-Yi University of Technology, CN=*.ncut.edu.tw"),
            ("Issuer", "C=TW, O=TAIWAN-CA, CN=TWCA Secure SSL Certification Authority"),
            ("有效期間", "2025-10-09 07:21:17 GMT 至 2026-11-05 15:59:59 GMT"),
            ("Serial", "47E9000000088AB31B59C27CF5366E40"),
            ("SAN", "DNS:*.ncut.edu.tw, DNS:ncut.edu.tw"),
            ("服務名稱覆蓋", "eduroam.ncut.edu.tw 被 *.ncut.edu.tw 涵蓋。"),
        ],
        label_width=2200,
        value_width=7160,
    )

    add_heading(doc, "3.4 本機帳號檔摘要", 2)
    users = []
    for line in cap.get("LOCAL_USERS", "").splitlines():
        if ":" in line:
            line_no, rest = line.split(":", 1)
            users.append((line_no, rest.strip(), "密碼已遮蔽；正式維運不得在文件中保存明文密碼。"))
    add_table(doc, ["行號", "帳號設定摘要", "說明"], users, [900, 5600, 2860], font_size=8.2)


def account_sop(doc):
    add_heading(doc, "4. 帳號新增、修改與停用 SOP", 1)
    add_callout(
        doc,
        "安全原則",
        "不要把使用者密碼寫入 Word、Email、Teams/LINE 對話或命令歷史紀錄。"
        "\n若密碼含有 !、@、$、引號、反斜線等特殊字元，建議用 sudoedit/vi 互動式編輯，避免 shell quoting 造成錯誤。",
        fill=WARN_FILL,
    )

    add_heading(doc, "4.1 新增本校 eduroam 使用者", 2)
    add_steps(
        doc,
        [
            "SSH 登入 RADIUS 主機 140.128.71.249，確認目前服務正常：systemctl is-active radiusd。",
            "備份本機帳號檔：cp -a /etc/raddb/mods-config/files/authorize /etc/raddb/mods-config/files/authorize.codex-backup-$(date +%Y%m%d-%H%M%S)。",
            "使用 sudoedit 或 vi 編輯 /etc/raddb/mods-config/files/authorize。",
            "同時建立短帳號與完整 realm 帳號兩筆，讓 PEAP 內層身份與測試工具都能明確命中。",
            "執行 radiusd -XC 檢查設定。",
            "重啟服務：systemctl restart radiusd，並確認 active/enabled。",
            "用 PAP 與 MSCHAP 測試 user@ncut.edu.tw，確認 Access-Accept。",
        ],
    )
    add_code(
        doc,
        """
# /etc/raddb/mods-config/files/authorize
newuser Cleartext-Password := "REPLACE_WITH_STRONG_PASSWORD", Proxy-To-Realm := LOCAL
newuser@ncut.edu.tw Cleartext-Password := "REPLACE_WITH_STRONG_PASSWORD", Proxy-To-Realm := LOCAL
""",
        "新增帳號範例",
    )

    add_heading(doc, "4.2 修改既有使用者密碼", 2)
    add_steps(
        doc,
        [
            "先確認要修改的帳號是否同時有短帳號與完整帳號，例如 danny 與 danny@ncut.edu.tw。",
            "備份 authorize 檔。",
            "修改兩筆 Cleartext-Password，並保留 Proxy-To-Realm := LOCAL。",
            "執行 radiusd -XC。",
            "重啟 radiusd。",
            "用新密碼測試 Access-Accept，再用舊密碼確認本機已拒絕。",
        ],
    )
    add_code(
        doc,
        """
sudo grep -nE '^(danny|danny@ncut\\.edu\\.tw)[[:space:]]' /etc/raddb/mods-config/files/authorize
sudo radiusd -XC
sudo systemctl restart radiusd
sudo systemctl is-active radiusd
""",
        "修改後基本檢查",
    )

    add_heading(doc, "4.3 停用或移除使用者", 2)
    add_bullets(
        doc,
        [
            "短期停用：在對應帳號行前加 # 註解，保留歷史紀錄與回復依據。",
            "永久移除：備份後刪除短帳號與完整 realm 帳號兩筆，完成後執行 radiusd -XC 與重啟。",
            "若只刪除其中一筆，可能出現某些測試方式拒絕、某些 EAP 內層身份仍接受的混亂狀態。",
        ]
    )

    add_heading(doc, "4.4 帳號測試命令", 2)
    add_code(
        doc,
        """
# 取得 localhost client secret，僅在互動 shell 暫存，不寫入文件或命令紀錄。
secret=$(sudo awk '
  /^[[:space:]]*client[[:space:]]+localhost[[:space:]]*\\{/ {in_block=1; next}
  in_block && /^[[:space:]]*\\}/ {in_block=0}
  in_block && /^[[:space:]]*secret[[:space:]]*=/ {
    sub(/^[[:space:]]*secret[[:space:]]*=[[:space:]]*/, "")
    sub(/[[:space:]]*(#.*)?$/, "")
    print
    exit
  }
' /etc/raddb/clients.conf)

radtest -t pap    'user@ncut.edu.tw' 'REPLACE_WITH_PASSWORD' 127.0.0.1 0 "$secret"
radtest -t mschap 'user@ncut.edu.tw' 'REPLACE_WITH_PASSWORD' 127.0.0.1 0 "$secret"
""",
        "本機 PAP/MSCHAP 驗證",
    )
    add_callout(
        doc,
        "判讀標準",
        "Access-Accept 代表本機 FreeRADIUS 帳號密碼、realm 與 MSCHAP 基本可用。"
        "\n若本機 Access-Accept 但 Wi-Fi 無法加入，下一步查用戶端憑證信任、無線控制器是否送到本機、以及 /var/log/radius/radius.log 的 EAP 訊息。",
        fill=OK_FILL,
    )


def cert_sop(doc):
    add_heading(doc, "5. SSL 憑證更新 SOP", 1)
    add_callout(
        doc,
        "更新時機",
        "目前憑證有效至 2026-11-05 15:59:59 GMT。建議到期前 30-45 天完成新憑證申請、安裝與用戶端抽測。",
        fill=WARN_FILL,
    )

    add_heading(doc, "5.1 憑證檔案需求", 2)
    add_table(
        doc,
        ["檔案", "用途", "目前命名"],
        [
            ("Private key", "伺服器私鑰；不可外流。", "ncutserver.key"),
            ("Server certificate", "eduroam.ncut.edu.tw 所屬 SSL/TLS 憑證。", "server.cer -> ncutserver.cer"),
            ("Intermediate CA", "TWCA 中繼憑證。", "uca.cer -> ncutserver-uca.cer"),
            ("Full chain", "server certificate + intermediate CA 串接。", "ncutserver-fullchain.pem"),
        ],
        [2200, 3900, 3260],
        font_size=8.8,
    )

    add_heading(doc, "5.2 更新步驟", 2)
    add_steps(
        doc,
        [
            "將新 private key、server certificate、intermediate CA 上傳至 /home/admindanny 暫存目錄。",
            "使用 openssl 確認 private key 可讀、server certificate 與 key modulus 相同。",
            "確認憑證 SAN 至少包含 *.ncut.edu.tw 或 eduroam.ncut.edu.tw。",
            "驗證 server certificate 與 intermediate chain：openssl verify -untrusted intermediate server.cer。",
            "備份 /etc/raddb/mods-available/eap。",
            "用 install 安裝至 /etc/raddb/certs，權限 root:radiusd 0640。",
            "串接 fullchain，確認 EAP 指到新 key/fullchain/intermediate。",
            "執行 radiusd -XC，通過後重啟 radiusd。",
            "執行本機 PAP/MSCHAP 測試與實際 Wi-Fi 抽測。",
            "刪除 /home/admindanny 暫存私鑰檔。",
        ],
    )
    add_code(
        doc,
        """
cd /home/admindanny

openssl x509 -in server.cer -noout -subject -issuer -dates -serial -ext subjectAltName
openssl x509 -in uca.cer -noout -subject -issuer -dates -serial
openssl rsa -in ncutserver.key -check -noout

cert_md5=$(openssl x509 -in server.cer -noout -modulus | openssl md5 | awk '{print $2}')
key_md5=$(openssl rsa -in ncutserver.key -noout -modulus | openssl md5 | awk '{print $2}')
[ "$cert_md5" = "$key_md5" ] && echo 'key/cert match'

openssl verify -untrusted uca.cer server.cer
""",
        "憑證與私鑰配對檢查",
    )
    add_code(
        doc,
        """
ts=$(date +%Y%m%d-%H%M%S)
cp -a /etc/raddb/mods-available/eap /etc/raddb/mods-available/eap.backup-${ts}

install -o root -g radiusd -m 0640 ncutserver.key /etc/raddb/certs/ncutserver.key
install -o root -g radiusd -m 0640 server.cer /etc/raddb/certs/ncutserver.cer
install -o root -g radiusd -m 0640 uca.cer /etc/raddb/certs/ncutserver-uca.cer
cat /etc/raddb/certs/ncutserver.cer /etc/raddb/certs/ncutserver-uca.cer \\
  > /etc/raddb/certs/ncutserver-fullchain.pem
chown root:radiusd /etc/raddb/certs/ncutserver-fullchain.pem
chmod 0640 /etc/raddb/certs/ncutserver-fullchain.pem
restorecon -v /etc/raddb/certs/ncutserver* 2>/dev/null || true
""",
        "安裝憑證檔",
    )
    add_code(
        doc,
        """
# /etc/raddb/mods-enabled/eap
private_key_file = ${certdir}/ncutserver.key
certificate_file = ${certdir}/ncutserver-fullchain.pem
ca_file = ${certdir}/ncutserver-uca.cer

radiusd -XC
systemctl restart radiusd
systemctl is-active radiusd
""",
        "EAP 設定與重啟",
    )

    add_heading(doc, "5.3 更新後驗證", 2)
    add_table(
        doc,
        ["驗證項目", "命令 / 方法", "通過條件"],
        [
            ("設定語法", "radiusd -XC", "最後顯示 configuration appears to be OK 或命令 exit code 0。"),
            ("服務狀態", "systemctl is-active radiusd", "active。"),
            ("Port", "ss -lunp | grep -E ':(1812|1813)'", "1812/udp 與 1813/udp 皆由 radiusd 監聽。"),
            ("憑證名稱", "openssl x509 -in /etc/raddb/certs/ncutserver-fullchain.pem -noout -subject -ext subjectAltName", "SAN 涵蓋 eduroam.ncut.edu.tw。"),
            ("帳號驗證", "radtest -t mschap user@ncut.edu.tw ...", "Access-Accept。"),
            ("實機 Wi-Fi", "手機/筆電忘記 eduroam 後重連", "可加入網路，伺服器名稱 eduroam.ncut.edu.tw。"),
        ],
        [2100, 4200, 3060],
        font_size=8.4,
    )


def operations_and_troubleshooting(doc):
    add_heading(doc, "6. 日常檢查與測試", 1)
    add_code(
        doc,
        """
hostname
ip -4 addr show
systemctl status radiusd --no-pager
systemctl status firewalld --no-pager
systemctl status openvpn-client@client --no-pager
firewall-cmd --list-ports
ip route get 10.1.77.7
ip route get 10.1.77.9
tail -f /var/log/radius/radius.log
""",
        "每日或異常時快速檢查",
    )

    add_heading(doc, "6.1 故障排除對照表", 2)
    add_table(
        doc,
        ["現象", "優先檢查", "可能處置"],
        [
            (
                "ncut@rc.edu.tw 可用，但 danny@ncut.edu.tw 不可用",
                "檢查 /etc/raddb/proxy.conf 是否有 realm ncut.edu.tw local；檢查 users 檔是否有完整帳號。",
                "新增或修正 local realm；確認 user@ncut.edu.tw 以 MSCHAP Access-Accept。",
            ),
            (
                "本機 radtest Access-Accept，但 Wi-Fi 仍無法加入",
                "檢查用戶端是否信任 TWCA chain、RADIUS Server Name 是否填 eduroam.ncut.edu.tw；查看 radius.log EAP 訊息。",
                "讓用戶端忘記 eduroam 後重設；必要時抓 radiusd debug 或控制器 log。",
            ),
            (
                "外校帳號大量 Home Server failed to respond",
                "檢查 tun0、openvpn-client@client、ip route get 10.1.77.7/10.1.77.9。",
                "重啟 OpenVPN；確認防火牆或 TANRC 端 10.1.77.7/10.1.77.9 可達。",
            ),
            (
                "radiusd 無法啟動",
                "執行 radiusd -XC 看語法錯誤；確認 eap 憑證路徑與權限。",
                "依錯誤行號修正設定；必要時用最近的 .codex-backup 還原。",
            ),
            (
                "更換憑證後用戶端跳憑證錯誤",
                "檢查 SAN、issuer、fullchain、ca_file；確認用戶端伺服器名稱。",
                "修正 fullchain 與 ca_file，重新發佈用戶端設定建議。",
            ),
            (
                "只出現 Access-Reject，沒有本機帳號 log",
                "檢查 request 是否進到本機；client 是否在 clients.conf；realm 是否被 proxy。",
                "補 clients.conf 或 realm；重啟 radiusd 後再測。",
            ),
        ],
        [2500, 3600, 3260],
        font_size=8.1,
    )

    add_heading(doc, "6.2 Log 查詢", 2)
    add_code(
        doc,
        """
# 最近錯誤
sudo tail -300 /var/log/radius/radius.log | \\
  grep -Ei 'Login incorrect|Home Server failed|No EAP session|Access-Reject|Access-Accept'

# 查特定帳號，注意輸出可能含個資，勿貼到公開文件
sudo grep -Ei 'danny|user@ncut\\.edu\\.tw' /var/log/radius/radius.log | tail -80
""",
        "常用 log 過濾",
    )

    add_heading(doc, "6.3 備份與回復", 2)
    add_bullets(
        doc,
        [
            "每次修改 /etc/raddb/proxy.conf、/etc/raddb/mods-available/eap、/etc/raddb/mods-config/files/authorize 前先 cp -a 備份。",
            "備份命名建議加時間戳，例如 proxy.conf.backup-YYYYMMDD-HHMMSS。",
            "回復時先複製原備份覆蓋目標，再執行 radiusd -XC；確認通過後才 restart radiusd。",
            "不要批次刪除備份檔；若空間不足，應逐一確認後單檔刪除。",
        ]
    )


def firewall_section(doc):
    add_heading(doc, "7. 防火牆與網路規則", 1)
    add_table(
        doc,
        ["方向", "來源", "目的", "Port / Protocol", "用途"],
        [
            ("Inbound", "校內無線控制器 / TANRC 節點", "140.128.71.249", "1812/udp", "RADIUS authentication。"),
            ("Inbound", "校內無線控制器 / TANRC 節點", "140.128.71.249", "1813/udp", "RADIUS accounting。"),
            ("Inbound 管理", "120.108.9.6 或核准管理 IP", "140.128.71.249", "22/tcp", "SSH 維運。"),
            ("Outbound", "140.128.71.249 / tun0", "10.1.77.7, 10.1.77.9", "1812/udp", "轉送外校 roaming authentication。"),
            ("Outbound", "140.128.71.249", "TANRC OpenVPN endpoint", "443/tcp 或實際 VPN 設定", "建立 OpenVPN tunnel。"),
            ("DNS", "140.128.71.249", "校內 DNS resolver", "53/udp,tcp", "解析 eduroam 與套件/更新需求。"),
        ],
        [1600, 2500, 2100, 1700, 1460],
        font_size=8.0,
    )
    add_callout(
        doc,
        "目前主機 firewalld",
        "firewalld 已啟用，public zone 目前開放 1812/udp、1813/udp。"
        "\n邊界防火牆仍須允許校內控制器與 TANRC 節點到 140.128.71.249 的 RADIUS 流量。",
        fill=OK_FILL,
    )


def change_log(doc):
    add_heading(doc, "8. 近期變更紀錄", 1)
    add_table(
        doc,
        ["日期", "變更", "驗證"],
        [
            ("2026-06-25", "套用 *.ncut.edu.tw SSL 憑證至 FreeRADIUS EAP。", "radiusd -XC OK；服務 active；本機 RADIUS Access-Accept。"),
            ("2026-06-25", "新增 realm ncut.edu.tw 與 eduroam.ncut.edu.tw 為本機處理。", "本校 danny@ncut.edu.tw PAP/MSCHAP Access-Accept；ncut@rc.edu.tw proxy Access-Accept。"),
            ("2026-06-25", "確認 eduroam.ncut.edu.tw DNS 指向 140.128.71.249。", "SAN DNS:*.ncut.edu.tw 可涵蓋。"),
            ("2026-06-25", "建立/調整本機測試帳號。", "密碼不列入文件；帳號檔摘要已遮蔽。"),
        ],
        [1600, 4700, 3060],
        font_size=8.5,
    )
    add_heading(doc, "8.1 已知注意事項", 2)
    add_bullets(
        doc,
        [
            "目前 log 仍可見部分外校或子網域帳號出現 Home Server failed to respond，需依實際來源判斷是否為 TANRC 端、VPN 或 realm 政策問題。",
            "直接對 TANRC 節點 radtest 可能不代表本機即時密碼狀態；帳號密碼修改後應優先以本機 127.0.0.1 PAP/MSCHAP 驗證。",
            "EAP 實機連線問題通常同時牽涉用戶端憑證信任、無線控制器轉送、realm 判斷與 inner-tunnel，不宜只看單一 Access-Reject。",
        ]
    )


def appendix(doc):
    add_heading(doc, "附錄 A：命令速查", 1)
    add_code(
        doc,
        """
# FreeRADIUS syntax check
sudo radiusd -XC

# Restart and status
sudo systemctl restart radiusd
sudo systemctl is-active radiusd
sudo systemctl is-enabled radiusd

# Firewall
sudo firewall-cmd --list-ports
sudo firewall-cmd --add-port=1812/udp --permanent
sudo firewall-cmd --add-port=1813/udp --permanent
sudo firewall-cmd --reload

# VPN / route
ip -4 addr show tun0
ip route get 10.1.77.7
ip route get 10.1.77.9
sudo systemctl status openvpn-client@client --no-pager

# Certificate
openssl x509 -in /etc/raddb/certs/ncutserver-fullchain.pem \\
  -noout -subject -issuer -dates -serial -ext subjectAltName
""",
        "常用維運命令",
    )

    doc.add_page_break()
    add_heading(doc, "附錄 B：目前本機 realm 設定片段", 1)
    add_code(
        doc,
        """
realm LOCAL {
}

realm ncut.edu.tw {
        nostrip
}

realm eduroam.ncut.edu.tw {
        nostrip
}

realm DEFAULT {
        pool = TANRC_POOL
        nostrip
}
""",
        "proxy.conf 摘要",
    )


def build():
    cap = parse_capture(CAPTURE)
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc)
    title_page(doc, cap)
    system_overview(doc, cap)
    config_inventory(doc, cap)
    current_radius_config(doc, cap)
    account_sop(doc)
    cert_sop(doc)
    operations_and_troubleshooting(doc)
    firewall_section(doc)
    change_log(doc)
    appendix(doc)

    core = doc.core_properties
    core.title = "NCUT eduroam RADIUS 技術說明文件"
    core.subject = "FreeRADIUS on Rocky Linux 9"
    core.author = "NCUT / Codex"
    core.keywords = "eduroam, RADIUS, FreeRADIUS, Rocky Linux 9, SSL, SOP"
    core.comments = "Sensitive values redacted."
    core.last_modified_by = "Codex"
    core.created = datetime(2026, 6, 25, 15, 40, 0)
    core.modified = datetime(2026, 6, 25, 15, 40, 0)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
